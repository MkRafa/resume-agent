"""File -> Document. Typed text and uploads converge on the same shape.

Fallback ladder: native text layer -> (future) OCR -> ask the user to paste.
No LLM is involved in getting bytes into text; that keeps extraction cheap,
deterministic, and - when it matters - entirely local.
"""

from __future__ import annotations

from pathlib import Path

from app.schemas import Document

TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".heic"}


class UnsupportedDocument(ValueError):
    pass


def from_text(text: str) -> Document:
    return Document(
        source_type="typed",
        raw_text=text.strip(),
        extraction_method="typed",
        confidence=1.0,
    )


def _from_pdf(path: Path) -> Document:
    try:
        import pymupdf
    except ImportError:  # pragma: no cover - older wheels only expose `fitz`
        try:
            import fitz as pymupdf  # type: ignore[no-redef]
        except ImportError as exc:
            raise UnsupportedDocument(
                "PDF support needs pymupdf: pip install pymupdf"
            ) from exc

    doc = pymupdf.open(path)
    try:
        chunks = [page.get_text("text") for page in doc]
        pages = doc.page_count
    finally:
        doc.close()

    text = "\n".join(chunks).strip()
    warnings: list[str] = []
    confidence = 1.0

    # A PDF with a near-empty text layer is a scan. We surface it rather than
    # silently handing 40 characters to the extractor and producing a hollow graph.
    if len(text) < 200:
        warnings.append(
            "PDF has little or no text layer - it is probably a scan. "
            "Run OCR or paste the text instead."
        )
        confidence = 0.2

    return Document(
        source_type="pdf",
        raw_text=text,
        filename=path.name,
        pages=pages,
        extraction_method="pymupdf-textlayer",
        confidence=confidence,
        warnings=warnings,
    )


def _from_docx(path: Path) -> Document:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedDocument(
            "DOCX support needs python-docx: pip install python-docx"
        ) from exc

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    # Resume content frequently hides in tables; skipping them loses whole roles.
    for table in d.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)

    return Document(
        source_type="docx",
        raw_text="\n".join(p for p in parts if p.strip()).strip(),
        filename=path.name,
        extraction_method="python-docx",
        confidence=1.0,
    )


def from_file(path: str | Path) -> Document:
    p = Path(path).expanduser()
    if not p.exists():
        raise FileNotFoundError(f"No such file: {p}")

    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _from_pdf(p)
    if suffix in {".docx", ".dotx"}:
        return _from_docx(p)
    if suffix in TEXT_SUFFIXES:
        return Document(
            source_type="md" if suffix != ".txt" else "txt",
            raw_text=p.read_text(encoding="utf-8", errors="replace").strip(),
            filename=p.name,
            extraction_method="plaintext",
            confidence=1.0,
        )
    if suffix in IMAGE_SUFFIXES:
        # Images go to the multimodal model rather than a local OCR pipeline;
        # the extractor node handles this branch.
        return Document(
            source_type="image",
            raw_text="",
            filename=p.name,
            extraction_method="deferred-multimodal",
            confidence=0.7,
            warnings=["Image input is passed to the multimodal extractor."],
        )
    if suffix == ".doc":
        raise UnsupportedDocument(
            "Legacy .doc is not supported. Save as .docx or PDF and retry."
        )
    raise UnsupportedDocument(f"Unsupported file type: {suffix or '(none)'}")


def load_input(text: str | None = None, file: str | Path | None = None) -> Document:
    """Intake router. Accepts typed text, a file, or both (text wins as the
    authoritative source; the file is treated as supplementary)."""
    if text and text.strip():
        doc = from_text(text)
        if file:
            extra = from_file(file)
            doc.raw_text = f"{doc.raw_text}\n\n--- attached: {extra.filename} ---\n{extra.raw_text}"
            doc.warnings.extend(extra.warnings)
        return doc
    if file:
        return from_file(file)
    raise ValueError("Provide typed text, a file, or both.")
